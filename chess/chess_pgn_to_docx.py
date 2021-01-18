import chess
import chess.pgn
import chess.svg
import os

from docx import Document
from docx.shared import Inches, Cm
from cairosvg import svg2png

temp_svg_file = "chess_pgn_to_docx_temp_svg.svg"
temp_png_file = "chess_pgn_to_docx_temp_svg.png"

style_black_and_white = '''
.square.light { fill: #cacaca; }
.square.dark  { fill: #898989; }
.square.light.lastmove { fill: #c3d889; }
.square.dark.lastmove { fill: #92b167; }
.check { fill: url(#check_gradient); }
.arrow { stroke: #ff5858; fill: #ff5858; }
.mark { stroke: #959fff; fill: #959fff;}
'''

style_mono_print = '''
.square.light { fill: #ffffff; }
.square.dark  { fill: #cacaca; }
.square.light.lastmove { fill: #ffffff; stroke: #000000;  stroke-width: 5;  stroke-linecap: butt;  stroke-dasharray: 0; }
.square.dark.lastmove { fill: #797979; stroke: #000000;  stroke-width: 5;  stroke-linecap: butt;  stroke-dasharray: 0; }
.check { fill: url(#check_gradient); }
.arrow { stroke: #ff5858; fill: #ff5858; }
.mark { stroke: #959fff; fill: #959fff;}
'''

def set_number_of_columns(section, cols):
    """ sets number of columns through xpath. """
    # # https://github.com/python-openxml/python-docx/issues/167
    WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))
    # The 'sep' option is similar with 'num'.
    WNS_COLS_SEP = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sep"
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_SEP, "1")

def set_margin(section):
    # changing the page margins => https://stackoverflow.com/questions/32914595/modify-docx-page-margins-with-python-docx
    # section.top_margin = Cm(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def update_docx_setting(document):
    sections = document.sections
    for section in sections:
        set_margin(section)
        set_number_of_columns(section,3)


def write_list_number(document, list_line):
    p = document.add_paragraph("", style='List Number')
    p.add_run(list_line).bold = True

def clean_temp_files():
    os.remove(temp_svg_file)
    os.remove(temp_png_file)

def write_board_screenshot(document, board_in_svg):
    with open(workspace + temp_svg_file, 'w') as content_file:
        content_file.write(board_in_svg)  # create svg file.
    with open(workspace + temp_svg_file, 'r') as content_file:
        content = content_file.read()
        svg2png(bytestring=content, write_to=temp_png_file)  # convert to png.
    document.add_picture(temp_png_file, width=Inches(2))  # TODO: to show 4(row) * 3(column) boards, 2.2 is the max size.

def write_title(document, title):
    document.add_heading(title, 0)

def is_white_move(step, first_move_white):
    # If it is a game started from the beginning, the steps 1,3,5 are white.
    if step % 2 == 1:
        return first_move_white
    else:
        return not first_move_white

def pgn_to_docx(workspace, pgn_filename, docx_filename,svg_style=None, flipped=False, show_initial_board=False, first_move_white=True):
    document = Document()
    update_docx_setting(document)
    write_title(document, pgn_filename)

    pgn = open(workspace + pgn_filename, encoding="utf-8")
    first_game = chess.pgn.read_game(pgn)
    print("Game Inforamtion:\n" + str(first_game) + "\n")

    comment_map = {}
    san_map = {}
    last_node = first_game.root()
    idx = 0

    while last_node.variations:
        # TODO: for variations.
        idx = idx + 1
        last_node = last_node.variations[0]
        comment_map[idx] = last_node.comment
        san_map[idx] = str(last_node.san())
        print("# " + str(idx))
        print("san: " + str(last_node.san())) #Comparing to last_node.move or last_node.uci(), we are familiar with SAN.
        print("variations: " + str(last_node.variations))
        print("comment: " + last_node.comment)
        if len(last_node.variations) > 1:
            another_node = last_node.variations[1]
            print("Other variations: " + str(another_node.variations))
        print("")

    board = first_game.board()
    if show_initial_board:
        write_list_number(document, "Initial Board")
        board_in_svg = chess.svg.board(board=board, flipped=flipped, size=400, style=svg_style)
        write_board_screenshot(document, board_in_svg)

    index = 1
    white = is_white_move(1, first_move_white) # set to the initial value of the first move.
    step = 0
    for move in first_game.mainline_moves():
        step+=1
        board.push(move)
        board_in_svg = chess.svg.board(board=board, lastmove=move, flipped=flipped, size=400, style=svg_style)

        turn = ""
        old_index = index
        old_white = white

        if white:
            turn = "White"
            white = False # change it for next loop
        else:
            turn = "Black"
            white = True # change it for next loop
            index += 1

        # write a line number item (e.g. "1. e4, e5") if it is a white move, or it is the first move.
        if old_white or (step == 1 and (not first_move_white)):
            if step == 1 and (not first_move_white):
                first_step = san_map[step]
                write_list_number(document, "... , " + first_step)
            else:
                write_step = san_map[step]
                blank_step = san_map[step + 1] if step + 1 in san_map else  "-"
                write_list_number(document, write_step + ", " + blank_step)

        # write the board screenshot and its comment if any.
        write_board_screenshot(document, board_in_svg)
        if comment_map[step]:
            document.add_paragraph(turn + " Comment: " + comment_map[step])


    print("Final status of this PGN file.")
    print(board)
    document.save(docx_filename)

    clean_temp_files()

if __name__ == '__main__':

    pgn_filename = "test_pgn_first_move_is_white.pgn"

    flipped = False
    show_initial_board = True
    first_move_white = True

    workspace = "./"
    docx_filename = pgn_filename.replace(".pgn","") + ".docx"

    pgn_to_docx(workspace, pgn_filename, docx_filename, svg_style=style_mono_print, flipped=flipped, show_initial_board=True, first_move_white=first_move_white)
